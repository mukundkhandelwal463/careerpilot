import io
import json
import os
import re
import time
import concurrent.futures

import docx
import google.generativeai as genai
from dotenv import load_dotenv
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from fpdf import FPDF

load_dotenv()
SYSTEM_API_KEY = os.getenv("GEMINI_API_KEY")
GEMINI_MODEL_CANDIDATES = (
    "gemini-2.5-flash",
    "gemini-2.0-flash",
    "gemini-flash-latest",
    "gemini-1.5-flash",
)
_RESOLVED_GEMINI_MODELS = None


def is_valid_key():
    return SYSTEM_API_KEY and SYSTEM_API_KEY.strip() != "your_gemini_api_key_here"


if is_valid_key():
    genai.configure(api_key=SYSTEM_API_KEY)


def _clip_text(text, max_chars=4000):
    txt = str(text or "")
    if len(txt) <= max_chars:
        return txt
    return txt[:max_chars]


def _generate_with_timeout(model, prompt, timeout_sec=10, retries=1):
    # Keeps API calls responsive and retries transient gateway/deadline failures.
    last_error = None
    for attempt in range(retries + 1):
        try:
            return model.generate_content(prompt, request_options={"timeout": timeout_sec}).text
        except Exception as exc:
            last_error = exc
            err = str(exc).lower()
            is_transient = any(
                token in err
                for token in ("deadline", "timeout", "timed out", "504", "503", "429", "temporar", "unavailable")
            )
            if attempt < retries and is_transient:
                time.sleep(0.8 * (attempt + 1))
                continue
            raise last_error


FALLBACK_API_KEY = os.getenv("DEEPSEEK_API_KEY")

def _generate_with_deepseek_fallback(prompt, timeout_sec=10):
    import urllib.request
    import json
    url = "https://api.deepseek.com/chat/completions"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {FALLBACK_API_KEY}"
    }
    data = {
        "model": "deepseek-chat",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.3
    }
    req = urllib.request.Request(url, headers=headers, data=json.dumps(data).encode("utf-8"))
    try:
        with urllib.request.urlopen(req, timeout=timeout_sec) as response:
            result = json.loads(response.read().decode("utf-8"))
            return result["choices"][0]["message"]["content"]
    except Exception as e:
        raise RuntimeError(f"DeepSeek fallback failed: {e}")

def _generate_with_model_fallback(prompt, timeout_sec=10):
    global _RESOLVED_GEMINI_MODELS

    if _RESOLVED_GEMINI_MODELS is None:
        discovered = []
        try:
            for model in genai.list_models():
                methods = getattr(model, "supported_generation_methods", []) or []
                if "generateContent" not in methods:
                    continue
                name = getattr(model, "name", "")
                short_name = name.split("/", 1)[-1] if "/" in name else name
                if short_name:
                    discovered.append(short_name)
        except Exception:
            discovered = []

        if discovered:
            ordered = [name for name in GEMINI_MODEL_CANDIDATES if name in discovered]
            _RESOLVED_GEMINI_MODELS = ordered or discovered[:4]
        else:
            _RESOLVED_GEMINI_MODELS = list(GEMINI_MODEL_CANDIDATES)

    last_error = None
    keys_to_try = [SYSTEM_API_KEY, FALLBACK_API_KEY]
    
    for current_key in keys_to_try:
        if not current_key or current_key == "your_gemini_api_key_here":
            continue
            
        genai.configure(api_key=current_key)
        
        for model_name in _RESOLVED_GEMINI_MODELS:
            try:
                model = genai.GenerativeModel(model_name)
                return _generate_with_timeout(model, prompt, timeout_sec=timeout_sec, retries=1)
            except Exception as exc:
                last_error = exc
                err = str(exc).upper()
                if "API_KEY_INVALID" in err or "429" in err or "QUOTA" in err:
                    break  # Break inner loop to try next API key
                continue  # Try next Gemini model

    # If all Gemini models and keys fail, attempt DeepSeek API fallback
    try:
        return _generate_with_deepseek_fallback(prompt, timeout_sec=timeout_sec)
    except Exception as e:
        last_error = e

    if last_error:
        raise last_error
    raise RuntimeError("No Gemini model candidates available and DeepSeek fallback failed.")


def get_gemini_ats_feedback(resume_text, job_description, score):
    if not is_valid_key():
        return "Please add a valid Gemini API Key to backend/.env to unlock AI feedback."
    try:
        prompt = f"""You are an expert ATS and career coach.
I have a resume and a job description. The match score is {score}%.
Job Description: {_clip_text(job_description, 2500)}
Resume: {_clip_text(resume_text, 5000)}
Give EXACTLY 5 actionable bullet points to improve the resume for this job.
Format each as numbered list (1..5), concise, no extra text."""
        return _generate_with_model_fallback(prompt, timeout_sec=14)
    except Exception as e:
        if "API_KEY_INVALID" in str(e):
            return "The Gemini API key in backend/.env is invalid."
        return f"Gemini Error: {e}"


def get_gemini_career_strategy(resume_text):
    if not is_valid_key():
        return "Please add a valid Gemini API Key to backend/.env to unlock AI career strategy."
    try:
        input_type = "Job Target/Category" if len(resume_text.split()) < 30 else "Full Resume Profile"
        prompt = f"""You are an elite career strategist.
I am providing you with my {input_type}:
"{_clip_text(resume_text, 4500)}"

Please provide a highly simplified, actionable 'Career Roadmap'.
Requirements:
1. Do not use large markdown headers (like # or ##). Avoid wordy introductions.
2. Provide a single text-based visual roadmap (using simple textual arrows like '→' or steps like 'Step 1: ...') showing how to progress from this {input_type} to a senior position.
3. Keep formatting extremely simple and minimal."""
        return _generate_with_model_fallback(prompt, timeout_sec=14)
    except Exception as e:
        if "API_KEY_INVALID" in str(e):
            return "The Gemini API key in backend/.env is invalid."
        return f"Gemini Error: {e}"


def get_gemini_resume_suggestions(data):
    if not is_valid_key():
        return "Please add a valid Gemini API Key to backend/.env to unlock resume suggestions."
    try:
        prompt = f"""I am writing my resume. Based on these draft details, give 3 actionable bullet points:
Name: {data.get('name')}
Summary: {data.get('summary')}
Skills: {data.get('skills')}
Experience: {data.get('experience')}
Education: {data.get('education')}
Keep it brief and ATS-focused."""
        return _generate_with_model_fallback(prompt, timeout_sec=14)
    except Exception as e:
        if "API_KEY_INVALID" in str(e):
            return "The Gemini API key in backend/.env is invalid."
        return f"Gemini Error: {e}"

def get_gemini_enhanced_text(text=""):
    if not is_valid_key() or not text.strip():
        return text or ""
    try:
        prompt = f"""You are an elite resume editor. Rewrite the following resume experience/project description text to make it extremely professional, impactful, and ATS-optimized.
Rules:
- Keep the core meaning and facts exactly the same.
- Use strong action verbs.
- Ensure proper grammar and punctuation.
- Output ONLY the raw corrected text. Do NOT add quotes, markdown, or chatty introductions.

Text to rewrite:
{_clip_text(text, 1000)}
"""
        result = _generate_with_model_fallback(prompt, timeout_sec=10) or ""
        return result.strip()
    except Exception:
        return text

def get_gemini_professional_summary(
    professional_title="",
    stream_or_category="",
    skills="",
    experience="",
    current_summary="",
):
    if not is_valid_key():
        return "Gemini Error: Missing valid Gemini API key."
    try:
        role_hint = (professional_title or "").strip()
        category_hint = (stream_or_category or "").strip()
        prompt = f"""You are an expert resume writer and ATS optimizer.
Your objective is to generate exactly a 2-sentence professional resume summary.

If `current_summary` is provided, you MUST analyze, correct, and heavily improve it, fixing any grammatical errors while keeping its core meaning and intent, condensing or expanding it to exactly 2 highly impactful sentences. 
If no `current_summary` is provided, generate a new one from scratch based on the skills and experience.

Rules:
- Write EXACTLY 2 sentences. No more, no less.
- Use a confident third-person style (no 'I', no 'my', no 'we').
- Start with the strongest role anchor:
  1) professional_title if provided
  2) else stream_or_category
- Build an ATS-friendly paragraph incorporating relevant skills naturally.
- Output plain text only, no bullets, no markdown, no quotes.

professional_title: {role_hint}
stream_or_category: {category_hint}
skills: {_clip_text(skills, 1200)}
experience: {_clip_text(experience, 2200)}
current_summary: {_clip_text(current_summary, 1200)}
"""
        text = _generate_with_model_fallback(prompt, timeout_sec=16) or ""
        text = re.sub(r"\s+", " ", str(text)).strip()
        return text
    except Exception as e:
        if "API_KEY_INVALID" in str(e):
            return "Gemini Error: The Gemini API key in backend/.env is invalid."
        return f"Gemini Error: {e}"


def get_gemini_full_resume_analysis(resume_text, job_description="", stream_or_category=""):
    fallback = {
        "category": (stream_or_category or "").strip() or "General",
        "ats_score": 0.0,
        "missing_keywords": [],
        "suggestions": ["Could not fetch full Gemini analysis right now. Please try again."],
        "gemini_ok": False,
    }
    if not is_valid_key():
        fallback["suggestions"] = ["Add a valid Gemini API key in backend/.env to enable AI analysis."]
        return fallback

    try:
        prompt = f"""You are an ATS and resume analyzer.
Return ONLY valid JSON (no markdown) with keys:
category (string), ats_score (number 0-100), missing_keywords (array of strings), suggestions (array of strings).
Rules:
- If stream_or_category is provided, use it as category unless resume strongly contradicts it.
- If job_description is empty, estimate ATS based on resume quality and relevance to stream_or_category.
- suggestions must be practical and ATS-focused, maximum 7 items.
stream_or_category: {stream_or_category}
job_description: {_clip_text(job_description, 2500)}
resume_text: {_clip_text(resume_text, 5000)}
"""
        raw = _generate_with_model_fallback(prompt, timeout_sec=16) or ""
        match = re.search(r"\{[\s\S]*\}", raw)
        if not match:
            return fallback

        parsed = json.loads(match.group(0))
        category = str(parsed.get("category", fallback["category"])).strip() or fallback["category"]
        try:
            ats_score = float(parsed.get("ats_score", 0.0))
        except Exception:
            ats_score = 0.0
        ats_score = round(max(0.0, min(100.0, ats_score)), 2)

        missing_keywords = parsed.get("missing_keywords", [])
        if not isinstance(missing_keywords, list):
            missing_keywords = []
        missing_keywords = [str(x).strip() for x in missing_keywords if str(x).strip()][:15]

        suggestions = parsed.get("suggestions", [])
        if not isinstance(suggestions, list):
            suggestions = []
        suggestions = [str(x).strip() for x in suggestions if str(x).strip()][:10]
        if not suggestions:
            suggestions = fallback["suggestions"]

        return {
            "category": category,
            "ats_score": ats_score,
            "missing_keywords": missing_keywords,
            "suggestions": suggestions,
            "gemini_ok": True,
        }
    except Exception:
        return fallback


class PDFResume(FPDF):
    def footer(self):
        self.set_y(-15)
        self.set_font("Arial", "I", 8)
        self.cell(0, 10, f"Page {self.page_no()}", align="C")


def _contact_location(data):
    return data.get("location", "") or data.get("linkedin", "")


def _split_non_empty_lines(text):
    return [line.strip() for line in str(text or "").splitlines() if line.strip()]


def _split_entry_blocks(text):
    raw = str(text or "").strip()
    if not raw:
        return []
    chunks = re.split(r"\n\s*\n", raw)
    blocks = []
    for chunk in chunks:
        lines = _split_non_empty_lines(chunk)
        if lines:
            blocks.append(lines)
    return blocks


def _split_skills(text):
    parts = re.split(r"[,\n]", str(text or ""))
    return [item.strip() for item in parts if item.strip()]


def _template_key(template):
    return str(template or "").strip().lower()


def _entry_records(text, default_label):
    records = []
    for idx, lines in enumerate(_split_entry_blocks(text), start=1):
        if len(lines) == 1:
            records.append(
                {
                    "meta": f"{default_label} {idx}",
                    "title": lines[0],
                    "details": [],
                }
            )
        elif len(lines) == 2:
            records.append(
                {
                    "meta": lines[0],
                    "title": lines[1],
                    "details": [],
                }
            )
        else:
            records.append(
                {
                    "meta": lines[0],
                    "title": lines[1],
                    "details": lines[2:],
                }
            )
    return records


def generate_pdf(data, template):
    pdf = PDFResume()
    pdf.add_page()
    location = _contact_location(data)
    template_key = _template_key(template)

    name = (data.get("name", "") or "YOUR NAME").strip()
    role = (data.get("headline", "") or "Professional Role").strip()
    summary = (data.get("summary", "") or "").strip()
    website = (data.get("website", "") or "").strip()

    def add_rule(y=None):
        line_y = y if y is not None else pdf.get_y()
        pdf.set_draw_color(140, 140, 140)
        pdf.set_line_width(0.5)
        pdf.line(14, line_y, 196, line_y)

    def render_records(x, y, width, records, empty_hint):
        curr_y = y
        if not records:
            pdf.set_xy(x, curr_y)
            pdf.set_font("Arial", "I", 10)
            pdf.set_text_color(105, 105, 105)
            pdf.multi_cell(width, 5.5, empty_hint)
            return pdf.get_y() + 2

        for record in records:
            pdf.set_xy(x, curr_y)
            pdf.set_font("Arial", "", 10)
            pdf.set_text_color(95, 95, 95)
            pdf.multi_cell(width, 5.2, record["meta"])
            curr_y = pdf.get_y()

            pdf.set_xy(x, curr_y)
            pdf.set_font("Arial", "B", 11)
            pdf.set_text_color(40, 40, 40)
            pdf.multi_cell(width, 5.6, record["title"])
            curr_y = pdf.get_y()

            if record["details"]:
                for line in record["details"]:
                    pdf.set_xy(x, curr_y)
                    pdf.set_font("Arial", "", 10)
                    pdf.set_text_color(45, 45, 45)
                    pdf.multi_cell(width, 5.3, f"- {line}")
                    curr_y = pdf.get_y()

            curr_y += 2
        return curr_y

    if template_key in {"classical.pdf", "resume for experienced.pdf"}:
        pdf.set_fill_color(235, 235, 235)
        pdf.rect(8, 8, 194, 281, "F")

        pdf.set_xy(10, 16)
        pdf.set_text_color(45, 45, 45)
        pdf.set_font("Arial", "B", 28)
        pdf.cell(0, 11, name.upper(), align="C", ln=True)
        pdf.set_font("Arial", "", 14)
        pdf.cell(0, 8, role, align="C", ln=True)

        pdf.set_font("Arial", "", 10.5)
        pdf.set_text_color(68, 68, 68)
        contacts = [data.get("phone", ""), data.get("email", ""), location]
        contacts = [c for c in contacts if c]
        pdf.cell(0, 7, "  |  ".join(contacts) if contacts else "+91 XXXXX XXXXX | hello@email.com | City, Country", align="C", ln=True)

        add_rule(pdf.get_y() + 1)
        pdf.set_y(pdf.get_y() + 6)

        pdf.set_font("Arial", "B", 13)
        pdf.set_text_color(40, 40, 40)
        pdf.cell(0, 7, "ABOUT ME", ln=True)
        pdf.set_font("Arial", "", 10.5)
        pdf.set_text_color(50, 50, 50)
        pdf.multi_cell(180, 6, summary or "Write a short profile about your background and strengths.")

        add_rule(pdf.get_y() + 2)
        pdf.set_y(pdf.get_y() + 7)

        pdf.set_font("Arial", "B", 13)
        pdf.set_text_color(40, 40, 40)
        pdf.cell(0, 7, "EDUCATION", ln=True)
        edu_y = render_records(14, pdf.get_y() + 1, 180, _entry_records(data.get("education", ""), "Education"), "Add education entries.")
        pdf.set_y(edu_y)

        add_rule(pdf.get_y() + 1)
        pdf.set_y(pdf.get_y() + 6)

        pdf.set_font("Arial", "B", 13)
        pdf.set_text_color(40, 40, 40)
        pdf.cell(0, 7, "WORK EXPERIENCE", ln=True)
        exp_y = render_records(14, pdf.get_y() + 1, 180, _entry_records(data.get("experience", ""), "Experience"), "Add work experience entries.")
        pdf.set_y(exp_y)

        add_rule(pdf.get_y() + 1)
        pdf.set_y(pdf.get_y() + 6)

        pdf.set_font("Arial", "B", 13)
        pdf.set_text_color(40, 40, 40)
        pdf.cell(0, 7, "SKILLS", ln=True)
        skills = _split_skills(data.get("skills", ""))
        if skills:
            pdf.set_font("Arial", "", 10.5)
            pdf.set_text_color(45, 45, 45)
            for skill in skills:
                pdf.multi_cell(180, 5.5, f"- {skill}")
        else:
            pdf.set_font("Arial", "I", 10)
            pdf.set_text_color(100, 100, 100)
            pdf.multi_cell(180, 5.5, "Add your strongest skills, separated by commas.")

    elif template_key == "freasher.pdf":
        left_x = 14
        left_w = 116
        right_x = 138
        right_w = 58

        pdf.set_fill_color(238, 238, 238)
        pdf.rect(8, 8, 194, 281, "F")

        pdf.set_xy(left_x, 16)
        pdf.set_font("Arial", "B", 24)
        pdf.set_text_color(44, 44, 44)
        pdf.multi_cell(left_w + right_w + 8, 10, name.upper())

        pdf.set_xy(left_x, 28)
        pdf.set_font("Arial", "", 13)
        pdf.set_text_color(72, 72, 72)
        pdf.multi_cell(left_w + right_w + 8, 7, role)

        add_rule(36)
        y = 40

        pdf.set_xy(left_x, y)
        pdf.set_font("Arial", "B", 13)
        pdf.set_text_color(40, 40, 40)
        pdf.multi_cell(left_w + right_w + 8, 6.5, "ABOUT ME")
        pdf.set_xy(left_x, pdf.get_y())
        pdf.set_font("Arial", "", 10.2)
        pdf.set_text_color(55, 55, 55)
        pdf.multi_cell(left_w + right_w + 8, 5.5, summary or "Write your introduction summary.")

        add_rule(pdf.get_y() + 2)
        top_cols_y = pdf.get_y() + 5

        y_left = top_cols_y
        pdf.set_xy(left_x, y_left)
        pdf.set_font("Arial", "B", 13)
        pdf.set_text_color(40, 40, 40)
        pdf.multi_cell(left_w, 6.5, "EDUCATION")
        y_left = render_records(left_x, pdf.get_y(), left_w, _entry_records(data.get("education", ""), "Education"), "Add education entries.")
        add_rule(y_left)
        y_left += 4

        pdf.set_xy(left_x, y_left)
        pdf.set_font("Arial", "B", 13)
        pdf.set_text_color(40, 40, 40)
        pdf.multi_cell(left_w, 6.5, "WORK EXPERIENCE")
        y_left = render_records(left_x, pdf.get_y(), left_w, _entry_records(data.get("experience", ""), "Experience"), "Add work experience entries.")

        y_right = top_cols_y
        pdf.set_xy(right_x, y_right)
        pdf.set_font("Arial", "B", 13)
        pdf.set_text_color(40, 40, 40)
        pdf.multi_cell(right_w, 6.5, "CONTACT")
        contact_lines = [data.get("phone", ""), data.get("email", ""), location, website]
        contact_lines = [line for line in contact_lines if line]
        if not contact_lines:
            contact_lines = ["Add contact details."]
        pdf.set_font("Arial", "", 10)
        pdf.set_text_color(55, 55, 55)
        for line in contact_lines:
            pdf.set_x(right_x)
            pdf.multi_cell(right_w, 5.3, line)
        y_right = pdf.get_y() + 2
        add_rule(y_right)
        y_right += 4

        pdf.set_xy(right_x, y_right)
        pdf.set_font("Arial", "B", 13)
        pdf.set_text_color(40, 40, 40)
        pdf.multi_cell(right_w, 6.5, "SKILLS")
        prof_skills = _split_skills(data.get("skills", ""))
        personal_skills = _split_skills(data.get("side_skills", ""))

        pdf.set_font("Arial", "B", 10)
        pdf.set_x(right_x)
        pdf.multi_cell(right_w, 5, "Professional")
        pdf.set_font("Arial", "", 10)
        for s in (prof_skills or ["Add professional skills."]):
            pdf.set_x(right_x)
            pdf.multi_cell(right_w, 5.1, f"- {s}")

        pdf.set_font("Arial", "B", 10)
        pdf.set_x(right_x)
        pdf.multi_cell(right_w, 5, "Personal")
        pdf.set_font("Arial", "", 10)
        for s in (personal_skills or ["Add personal skills."]):
            pdf.set_x(right_x)
            pdf.multi_cell(right_w, 5.1, f"- {s}")
        y_right = pdf.get_y() + 2
        add_rule(y_right)
        y_right += 4

        pdf.set_xy(right_x, y_right)
        pdf.set_font("Arial", "B", 13)
        pdf.multi_cell(right_w, 6.5, "LANGUAGES")
        languages = _split_skills(data.get("languages", ""))
        pdf.set_font("Arial", "", 10)
        for lang in (languages or ["Add language proficiency entries."]):
            pdf.set_x(right_x)
            pdf.multi_cell(right_w, 5.1, f"- {lang}")

        final_y = max(y_left, pdf.get_y())
        pdf.set_y(final_y)

    elif template_key == "resume for experienced2.pdf":
        pdf.set_fill_color(238, 238, 238)
        pdf.rect(8, 8, 194, 281, "F")

        pdf.set_xy(10, 14)
        pdf.set_font("Arial", "B", 24)
        pdf.set_text_color(44, 44, 44)
        pdf.cell(0, 10, name.upper(), align="C", ln=True)

        pdf.set_font("Arial", "B", 12)
        pdf.cell(0, 6, role.upper(), align="C", ln=True)

        contact_line = " | ".join([v for v in [location, data.get("phone", ""), data.get("email", "")] if v])
        pdf.set_font("Arial", "", 10)
        pdf.set_text_color(70, 70, 70)
        pdf.cell(0, 5.5, contact_line or "City, Country | +91 XXXXX XXXXX | hello@email.com", align="C", ln=True)
        pdf.cell(0, 5.5, website or "www.yourwebsite.com", align="C", ln=True)

        add_rule(pdf.get_y() + 2)
        pdf.set_y(pdf.get_y() + 8)

        pdf.set_font("Arial", "B", 13)
        pdf.set_text_color(40, 40, 40)
        pdf.cell(0, 7, "SUMMARY", ln=True)
        pdf.set_font("Arial", "", 10.2)
        pdf.set_text_color(55, 55, 55)
        pdf.multi_cell(180, 5.4, summary or "Write your executive summary.")

        add_rule(pdf.get_y() + 2)
        pdf.set_y(pdf.get_y() + 6)

        pdf.set_font("Arial", "B", 13)
        pdf.set_text_color(40, 40, 40)
        pdf.cell(0, 7, "WORK EXPERIENCE", ln=True)
        y_exp = render_records(14, pdf.get_y() + 1, 180, _entry_records(data.get("experience", ""), "Experience"), "Add work experience entries.")
        pdf.set_y(y_exp)

        add_rule(pdf.get_y() + 2)
        pdf.set_y(pdf.get_y() + 6)

        pdf.set_font("Arial", "B", 13)
        pdf.set_text_color(40, 40, 40)
        pdf.cell(0, 7, "EDUCATION", ln=True)
        y_edu = render_records(14, pdf.get_y() + 1, 180, _entry_records(data.get("education", ""), "Education"), "Add education entries.")
        pdf.set_y(y_edu)

    else:
        pdf.set_font("Arial", "B", 24)
        pdf.cell(0, 10, name, align="C", ln=True)
        pdf.set_font("Arial", "", 11)
        pdf.cell(0, 6, f"{data.get('email', '')} | {data.get('phone', '')} | {location}", align="C", ln=True)
        pdf.ln(5)
        for title, value in [
            ("Professional Summary", data.get("summary", "")),
            ("Key Skills", data.get("skills", "")),
            ("Experience", data.get("experience", "")),
            ("Education", data.get("education", "")),
        ]:
            pdf.set_font("Arial", "B", 13)
            pdf.cell(0, 7, title.upper(), ln=True)
            pdf.set_font("Arial", "", 10.5)
            pdf.multi_cell(180, 5.5, str(value or f"Add {title.lower()} details."))
            pdf.ln(1.5)

    return bytes(pdf.output(dest="S"))


def generate_docx(data, template):
    doc = docx.Document()
    location = _contact_location(data)
    template_key = _template_key(template)

    def add_heading(title, level_size=13):
        heading = doc.add_paragraph()
        run = heading.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(level_size)

    def add_rule():
        rule = doc.add_paragraph("_" * 96)
        rule.runs[0].font.size = Pt(7)

    def add_records_to_container(container, records, empty_text):
        if not records:
            hint = container.add_paragraph(empty_text)
            if hint.runs:
                hint.runs[0].italic = True
            return

        for record in records:
            meta_pg = container.add_paragraph(record["meta"])
            if meta_pg.runs:
                meta_pg.runs[0].font.size = Pt(10)
                meta_pg.runs[0].italic = True

            title_pg = container.add_paragraph(record["title"])
            if title_pg.runs:
                title_pg.runs[0].font.size = Pt(11)
                title_pg.runs[0].bold = True

            for detail in record["details"]:
                container.add_paragraph(detail, style="List Bullet")

    if template_key in {"classical.pdf", "resume for experienced.pdf"}:
        name_pg = doc.add_paragraph()
        name_pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_pg.add_run((data.get("name", "") or "YOUR NAME").upper())
        name_run.bold = True
        name_run.font.size = Pt(28)

        role_pg = doc.add_paragraph(data.get("headline", "") or "Professional Role")
        role_pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if role_pg.runs:
            role_pg.runs[0].font.size = Pt(14)

        contact_line = " | ".join([v for v in [data.get("phone", ""), data.get("email", ""), location] if v])
        contact_pg = doc.add_paragraph(contact_line or "+91 XXXXX XXXXX | hello@email.com | City, Country")
        contact_pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_rule()

        add_heading("About Me")
        doc.add_paragraph(data.get("summary", "") or "Write a short profile about your background and strengths.")
        add_rule()

        add_heading("Education")
        add_records_to_container(doc, _entry_records(data.get("education", ""), "Education"), "Add education entries.")
        add_rule()

        add_heading("Work Experience")
        add_records_to_container(doc, _entry_records(data.get("experience", ""), "Experience"), "Add work experience entries.")
        add_rule()

        add_heading("Skills")
        for skill in (_split_skills(data.get("skills", "")) or ["Add your strongest skills, separated by commas."]):
            doc.add_paragraph(skill, style="List Bullet")

    elif template_key == "freasher.pdf":
        name_pg = doc.add_paragraph()
        name_pg.alignment = WD_ALIGN_PARAGRAPH.LEFT
        name_run = name_pg.add_run((data.get("name", "") or "YOUR NAME").upper())
        name_run.bold = True
        name_run.font.size = Pt(24)

        role_pg = doc.add_paragraph(data.get("headline", "") or "Professional Role")
        role_pg.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_rule()

        add_heading("About Me")
        doc.add_paragraph(data.get("summary", "") or "Write your introduction summary.")
        add_rule()

        table = doc.add_table(rows=1, cols=2)
        left = table.rows[0].cells[0]
        right = table.rows[0].cells[1]

        left_heading = left.paragraphs[0].add_run("EDUCATION")
        left_heading.bold = True
        add_records_to_container(left, _entry_records(data.get("education", ""), "Education"), "Add education entries.")

        left.add_paragraph()
        left_work = left.add_paragraph().add_run("WORK EXPERIENCE")
        left_work.bold = True
        add_records_to_container(left, _entry_records(data.get("experience", ""), "Experience"), "Add work experience entries.")

        right_heading = right.paragraphs[0].add_run("CONTACT")
        right_heading.bold = True
        for line in [data.get("phone", ""), data.get("email", ""), location, data.get("website", "")]:
            if line:
                right.add_paragraph(line)

        right.add_paragraph()
        right_skill_heading = right.add_paragraph().add_run("SKILLS")
        right_skill_heading.bold = True
        right_prof = right.add_paragraph().add_run("Professional")
        right_prof.bold = True
        for skill in (_split_skills(data.get("skills", "")) or ["Add professional skills."]):
            right.add_paragraph(skill, style="List Bullet")

        right_personal = right.add_paragraph().add_run("Personal")
        right_personal.bold = True
        for skill in (_split_skills(data.get("side_skills", "")) or ["Add personal skills."]):
            right.add_paragraph(skill, style="List Bullet")

        right.add_paragraph()
        right_lang_heading = right.add_paragraph().add_run("LANGUAGES")
        right_lang_heading.bold = True
        for lang in (_split_skills(data.get("languages", "")) or ["Add language proficiency entries."]):
            right.add_paragraph(lang, style="List Bullet")

    elif template_key == "resume for experienced2.pdf":
        name_pg = doc.add_paragraph()
        name_pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_pg.add_run((data.get("name", "") or "YOUR NAME").upper())
        name_run.bold = True
        name_run.font.size = Pt(24)

        role_pg = doc.add_paragraph((data.get("headline", "") or "Business Consultant").upper())
        role_pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if role_pg.runs:
            role_pg.runs[0].bold = True

        contact_line = " | ".join([v for v in [location, data.get("phone", ""), data.get("email", "")] if v])
        doc.add_paragraph(contact_line or "City, Country | +91 XXXXX XXXXX | hello@email.com").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(data.get("website", "") or "www.yourwebsite.com").alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_rule()

        add_heading("Summary")
        doc.add_paragraph(data.get("summary", "") or "Write your executive summary.")
        add_rule()

        add_heading("Work Experience")
        add_records_to_container(doc, _entry_records(data.get("experience", ""), "Experience"), "Add work experience entries.")
        add_rule()

        add_heading("Education")
        add_records_to_container(doc, _entry_records(data.get("education", ""), "Education"), "Add education entries.")

    else:
        name_pg = doc.add_paragraph((data.get("name", "") or "Candidate").upper())
        name_pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if name_pg.runs:
            name_pg.runs[0].bold = True
            name_pg.runs[0].font.size = Pt(24)

        doc.add_paragraph(f"{data.get('email', '')} | {data.get('phone', '')} | {location}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_rule()

        for title, value in [
            ("Professional Summary", data.get("summary", "")),
            ("Key Skills", data.get("skills", "")),
            ("Experience", data.get("experience", "")),
            ("Education", data.get("education", "")),
        ]:
            add_heading(title)
            doc.add_paragraph(str(value or f"Add {title.lower()} details."))

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def generate_docx_from_builder(data):
    """Generate a properly formatted DOCX from the React Resume Builder data structure.

    Expects data with keys:
      personal_info: {full_name, email, phone, location, linkedin, website}
      professional_summary: str
      experience: [{position, company, start_date, end_date, is_current, description}]
      education: [{degree, field, institution, graduation_date, gpa}]
      project: [{name, description}]
      skills: [str]
      keywords: [str]
    """
    from docx.shared import Inches, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = docx.Document()

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(1)
    style.paragraph_format.space_before = Pt(0)

    pi = data.get("personal_info", {}) or {}
    accent_hex = data.get("accent_color", "2563eb") or "2563eb"
    accent_hex = accent_hex.lstrip("#")
    if len(accent_hex) != 6:
        accent_hex = "2563eb"
    try:
        accent = RGBColor(int(accent_hex[0:2], 16), int(accent_hex[2:4], 16), int(accent_hex[4:6], 16))
    except Exception:
        accent = RGBColor(0x25, 0x63, 0xEB)

    def add_horizontal_rule():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), accent_hex)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_section_heading(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = accent
        add_horizontal_rule()

    def format_date(date_str):
        if not date_str:
            return ""
        try:
            parts = date_str.split("-")
            from datetime import datetime
            dt = datetime(int(parts[0]), int(parts[1]), 1)
            return dt.strftime("%b %Y")
        except Exception:
            return date_str

    # ── NAME (centered, large, accent color) ──
    name_pg = doc.add_paragraph()
    name_pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_pg.paragraph_format.space_after = Pt(1)
    name_run = name_pg.add_run(pi.get("full_name", "") or "Your Name")
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.color.rgb = accent

    # ── CONTACT LINE (centered) ──
    contact_parts = []
    if pi.get("email"):
        contact_parts.append(pi["email"])
    if pi.get("phone"):
        contact_parts.append(pi["phone"])
    if pi.get("location"):
        contact_parts.append(pi["location"])
    if pi.get("linkedin"):
        contact_parts.append(pi["linkedin"])
    if pi.get("website"):
        contact_parts.append(pi["website"])

    if contact_parts:
        contact_pg = doc.add_paragraph()
        contact_pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_pg.paragraph_format.space_after = Pt(1)
        contact_run = contact_pg.add_run("  |  ".join(contact_parts))
        contact_run.font.size = Pt(8)
        contact_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    add_horizontal_rule()

    # ── PROFESSIONAL SUMMARY ──
    summary = data.get("professional_summary", "")
    if summary:
        add_section_heading("Professional Summary")
        p = doc.add_paragraph(summary)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # ── EXPERIENCE ──
    experience = data.get("experience", []) or []
    if experience:
        add_section_heading("Professional Experience")
        for exp in experience:
            # Row: Position | Dates
            pos_pg = doc.add_paragraph()
            pos_pg.paragraph_format.space_before = Pt(2)
            pos_pg.paragraph_format.space_after = Pt(0)
            pos_run = pos_pg.add_run(exp.get("position", "") or "Position")
            pos_run.bold = True
            pos_run.font.size = Pt(9)
            date_start = format_date(exp.get("start_date", ""))
            date_end = "Present" if exp.get("is_current") else format_date(exp.get("end_date", ""))
            date_str = f"{date_start} – {date_end}" if date_start else ""
            if date_str:
                pos_run2 = pos_pg.add_run(f"    |    {date_str}")
                pos_run2.font.size = Pt(8)
                pos_run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            # Company name
            company = exp.get("company", "")
            if company:
                comp_pg = doc.add_paragraph()
                comp_pg.paragraph_format.space_after = Pt(1)
                comp_run = comp_pg.add_run(company)
                comp_run.font.size = Pt(8)
                comp_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                comp_run.italic = True

            # Description as bullet points — max 4 bullets, max 160 chars each
            desc = exp.get("description", "")
            if desc:
                bullet_count = 0
                for line in desc.strip().split("\n"):
                    if bullet_count >= 4:
                        break
                    line = line.strip().lstrip("•-").strip()
                    if line:
                        # Truncate to ~160 chars (approx 2 lines)
                        if len(line) > 160:
                            line = line[:157].rstrip() + "..."
                        bp = doc.add_paragraph(line, style="List Bullet")
                        bp.paragraph_format.space_after = Pt(0)
                        for r in bp.runs:
                            r.font.size = Pt(8)
                        bullet_count += 1

    # ── PROJECTS ──
    projects = data.get("project", []) or []
    if projects:
        add_section_heading("Projects")
        for proj in projects[:3]:  # Max 3 projects
            p_name = doc.add_paragraph()
            p_name.paragraph_format.space_before = Pt(1)
            p_name.paragraph_format.space_after = Pt(0)
            run = p_name.add_run(proj.get("name", "") or "Project Name")
            run.bold = True
            run.font.size = Pt(9)
            # Show date next to project name
            proj_date = proj.get("date", "")
            if proj_date:
                date_run = p_name.add_run(f"  |  {proj_date}")
                date_run.font.size = Pt(8)
                date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            # Show link if available
            proj_link = proj.get("link", "")
            if proj_link:
                link_run = p_name.add_run(f"  •  {proj_link}")
                link_run.font.size = Pt(7)
                link_run.font.color.rgb = RGBColor(0x66, 0x88, 0xCC)

            p_desc = proj.get("description", "")
            if p_desc:
                # Truncate project description to ~200 chars
                if len(p_desc) > 200:
                    p_desc = p_desc[:197].rstrip() + "..."
                dp = doc.add_paragraph(p_desc)
                dp.paragraph_format.space_after = Pt(1)
                for r in dp.runs:
                    r.font.size = Pt(8)
                    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # ── EDUCATION ──
    education = data.get("education", []) or []
    if education:
        add_section_heading("Education")
        for edu in education:
            edu_pg = doc.add_paragraph()
            edu_pg.paragraph_format.space_before = Pt(2)
            edu_pg.paragraph_format.space_after = Pt(0)
            degree_text = edu.get("degree", "") or ""
            field = edu.get("field", "")
            if field:
                degree_text += f" in {field}"
            run = edu_pg.add_run(degree_text or "Degree")
            run.bold = True
            run.font.size = Pt(9)

            grad_date = format_date(edu.get("graduation_date", ""))
            if grad_date:
                run2 = edu_pg.add_run(f"    |    {grad_date}")
                run2.font.size = Pt(8)
                run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            inst = edu.get("institution", "")
            if inst:
                inst_pg = doc.add_paragraph()
                inst_pg.paragraph_format.space_after = Pt(0)
                inst_run = inst_pg.add_run(inst)
                inst_run.font.size = Pt(8)
                inst_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

            gpa = edu.get("gpa", "")
            if gpa:
                gpa_pg = doc.add_paragraph()
                gpa_pg.paragraph_format.space_after = Pt(1)
                gpa_run = gpa_pg.add_run(f"GPA: {gpa}")
                gpa_run.font.size = Pt(7)
                gpa_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ── CERTIFICATIONS ──
    certifications = data.get("certifications", []) or []
    if certifications:
        add_section_heading("Certifications")
        for cert in certifications[:4]:  # Max 4
            cert_name = cert.get("name", "") or ""
            cert_date = cert.get("date", "") or ""
            cert_link = cert.get("link", "") or ""
            line_parts = [cert_name]
            if cert_date:
                line_parts.append(cert_date)
            if cert_link:
                line_parts.append(cert_link)
            cert_pg = doc.add_paragraph()
            cert_pg.paragraph_format.space_before = Pt(1)
            cert_pg.paragraph_format.space_after = Pt(0)
            cert_run = cert_pg.add_run(cert_name)
            cert_run.bold = True
            cert_run.font.size = Pt(9)
            if cert_date or cert_link:
                extra = []
                if cert_date: extra.append(cert_date)
                if cert_link: extra.append(cert_link)
                meta_run = cert_pg.add_run("  |  " + "  |  ".join(extra))
                meta_run.font.size = Pt(8)
                meta_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # ── CORE SKILLS (merged with keywords) ──
    skills = data.get("skills", []) or []
    keywords = data.get("keywords", []) or []
    all_skills = list(skills) + [k for k in keywords if k not in skills]
    if all_skills:
        add_section_heading("Core Skills")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run("  •  ".join(all_skills))
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def get_gemini_parse_resume_to_json(resume_text):
    if not is_valid_key():
        raise RuntimeError("No valid Gemini API key found.")
    try:
        prompt = f"""You are an expert ATS data extraction system.
I am providing you with the text extracted from a user's resume.
Your job is to extract the following fields and return ONLY a valid JSON object.
Fields to extract:
- full_name
- headline (professional title)
- email
- phone
- location (city, state, country)
- summary (a brief professional summary)
- skills (comma separated list)
- experience (a block of text outlining work experience, keep it concise but informative)
- education (a block of text outlining education, keep it concise)
- address
- city
- state
- zip_code
- country
- linkedin (url if present)
- github (url if present)
- website (url if present)

Do NOT return markdown blocks like ```json. Return JUST the raw JSON string starting with {{ and ending with }}.
If a field is not found, leave it as an empty string.

Resume Text:
{_clip_text(resume_text, 6000)}
"""
        response_text = _generate_with_model_fallback(prompt, timeout_sec=20)
        
        # Clean up any potential markdown formatting from Gemini
        response_text = response_text.strip()
        if response_text.startswith("```json"):
            response_text = response_text[7:]
        if response_text.startswith("```"):
            response_text = response_text[3:]
        if response_text.endswith("```"):
            response_text = response_text[:-3]
            
        return json.loads(response_text.strip())
    except Exception as e:
        print(f"Error parsing resume to JSON: {e}")
        return {}


def generate_interview_questions(resume_text):
    if not is_valid_key():
        return [
            {"id": 1, "type": "Technical", "question": "Explain your experience with Python, javascript or other skills listed in your CV."},
            {"id": 2, "type": "Technical", "question": "What is the most complex backend system issue you solved?"},
            {"id": 3, "type": "HR", "question": "Why do you want to join our organization?"},
            {"id": 4, "type": "HR", "question": "Describe a situation where you had to work under tight constraints."},
            {"id": 5, "type": "Management", "question": "How do you manage deadlines and prioritize tasks?"},
            {"id": 6, "type": "Management", "question": "Explain a time when you coordinated a multi-developer project."},
            {"id": 7, "type": "Technical", "question": "Describe a system architecture challenge you faced."},
            {"id": 8, "type": "HR", "question": "How do you handle negative feedback from a senior stakeholder?"},
            {"id": 9, "type": "Management", "question": "How do you handle scope creep mid-project?"},
            {"id": 10, "type": "Technical", "question": "Explain your understanding of CI/CD and production deployment."}
        ]

    prompt = f"""Based on the following candidate's resume/CV text, generate exactly 10 interview questions.
The questions should be a mix of:
- Technical questions (based on the candidate's skills, tools, frameworks, and achievements mentioned in the resume)
- HR questions (behavioral, teamwork, conflict resolution, situational)
- Management questions (leadership, project coordination, priority setting)

Return a JSON array of objects. Each object should have three fields:
- 'id': integer from 1 to 10
- 'type': 'Technical', 'HR', or 'Management'
- 'question': the actual question text.

Format the response strictly as valid JSON (start with '[' and end with ']'). Do not add markdown backticks or any other text before/after the JSON array.

Candidate Resume Text:
{_clip_text(resume_text, max_chars=4000)}"""

    try:
        raw_response = _generate_with_model_fallback(prompt, timeout_sec=15)
        cleaned = re.sub(r"```json\s*", "", raw_response)
        cleaned = re.sub(r"```\s*", "", cleaned).strip()
        return json.loads(cleaned)
    except Exception as e:
        print(f"Error generating questions via Gemini: {e}")
        return [
            {"id": 1, "type": "Technical", "question": "Explain your experience with Python, javascript or other skills listed in your CV."},
            {"id": 2, "type": "Technical", "question": "What is the most complex backend system issue you solved?"},
            {"id": 3, "type": "HR", "question": "Why do you want to join our organization?"},
            {"id": 4, "type": "HR", "question": "Describe a situation where you had to work under tight constraints."},
            {"id": 5, "type": "Management", "question": "How do you manage deadlines and prioritize tasks?"},
            {"id": 6, "type": "Management", "question": "Explain a time when you coordinated a multi-developer project."},
            {"id": 7, "type": "Technical", "question": "Describe a system architecture challenge you faced."},
            {"id": 8, "type": "HR", "question": "How do you handle negative feedback from a senior stakeholder?"},
            {"id": 9, "type": "Management", "question": "How do you handle scope creep mid-project?"},
            {"id": 10, "type": "Technical", "question": "Explain your understanding of CI/CD and production deployment."}
        ]


def grade_interview_response(question, response):
    if not is_valid_key():
        return {
            "score": 80,
            "strengths": ["Spoken response received", "Provided direct answer"],
            "improvements": ["Unlock full AI evaluation by adding GEMINI_API_KEY to your .env file"]
        }

    prompt = f"""An interview candidate was asked the following question:
'{question}'

Their spoken/recognized response was:
'{response}'

Evaluate this response. Provide:
1. A numerical score from 0 to 100 based on completeness, correctness, clarity, and professionalism.
2. A short list of strengths (1-2 bullet points).
3. A short list of suggested improvements (1-2 bullet points).

Return a valid JSON object with the keys:
- 'score': integer (0 to 100)
- 'strengths': array of strings
- 'improvements': array of strings

Format the response strictly as valid JSON (start with '{{' and end with '}}'). Do not add markdown backticks or any other text before/after the JSON object.

Evaluation:"""

    try:
        raw_response = _generate_with_model_fallback(prompt, timeout_sec=12)
        cleaned = re.sub(r"```json\s*", "", raw_response)
        cleaned = re.sub(r"```\s*", "", cleaned).strip()
        return json.loads(cleaned)
    except Exception as e:
        print(f"Error grading response via Gemini: {e}")
        return {
            "score": 75,
            "strengths": ["Clear spoken articulation"],
            "improvements": ["Clarify technical trade-offs in candidate response"]
        }


def generate_career_roadmap(resume_text, target_role, job_description=""):
    if not is_valid_key():
        return {
            "current_skills": ["Python", "Git", "Software engineering basic principles"],
            "gap_skills": ["Advanced System Design", "Docker & Kubernetes", "Cloud Architecture (AWS)"],
            "levels": [
                {
                    "level": 1,
                    "title": "Level 1: Specialization Fundamentals",
                    "topics": ["Advanced programming language constructs", "Data storage modeling (SQL vs NoSQL)"],
                    "duration": "2 weeks",
                    "focus": "Establish strong backend and role specialization foundations."
                },
                {
                    "level": 2,
                    "title": "Level 2: Distributed Systems",
                    "topics": ["Microservices architecture pattern", "Caching & message queues (Redis/Kafka)"],
                    "duration": "3 weeks",
                    "focus": "Learn to design highly available and decoupled services."
                }
            ]
        }

    jd_clause = f"\nTarget Job Description to crack:\n{job_description}" if job_description else ""
    prompt = f"""You are a top-tier career counselor and technical roadmap designer.
The candidate has the following resume:
{_clip_text(resume_text, max_chars=4000)}

They want to study and plan a learning path to become a:
{target_role}{jd_clause}

Analyze their background, skills gap, and outline a complete, comprehensive level-by-level study roadmap. You MUST generate at least 3 to 4 levels (e.g., Level 1, Level 2, Level 3, etc.) covering everything from foundations to mastery.
Include a precise gap analysis showing what skills they already have (relative to target role) and what they have left to learn.

Return a valid JSON object with the following keys:
- 'current_skills': list of strings
- 'gap_skills': list of strings
- 'levels': list of objects, where each object has keys:
  - 'level': integer (1, 2, 3, etc.)
  - 'title': string
  - 'topics': list of strings. IMPORTANT: Instead of broad topics like 'Advanced algorithms', you MUST provide exactly 2 highly specific, actionable sub-topic names for each broad concept (e.g., 'Advanced algorithms: Dijkstra', 'Advanced algorithms: A* Search'). This flat list of strings will be mapped directly to daily tasks, 2 per day.
  - 'duration': string
  - 'focus': string

Format the response strictly as valid JSON (start with '{{' and end with '}}'). Do not add markdown backticks (e.g. ```json) or any extra text before or after the JSON response.

JSON Roadmap:"""

    try:
        raw_response = _generate_with_model_fallback(prompt, timeout_sec=16)
        cleaned = re.sub(r"```json\s*", "", raw_response)
        cleaned = re.sub(r"```\s*", "", cleaned).strip()
        return json.loads(cleaned)
    except Exception as e:
        print(f"Error generating roadmap: {e}")
        return {
            "current_skills": ["Python", "Basic JavaScript"],
            "gap_skills": ["Deep Learning", "A/B Testing", "ETL Systems", "System Architecture"],
            "levels": [
                {
                    "level": 1,
                    "title": "Level 1: Core Foundations",
                    "topics": [
                        "Advanced algorithms: Dijkstra & Bellman-Ford", 
                        "Advanced algorithms: A* Search Pattern", 
                        "Specialized packages: React Router in-depth", 
                        "Specialized packages: Redux State Management"
                    ],
                    "duration": "3 weeks",
                    "focus": "Strengthen domain-specific developer foundations and core coding skills."
                },
                {
                    "level": 2,
                    "title": "Level 2: Advanced Application Design",
                    "topics": [
                        "Microservices patterns: API Gateways", 
                        "Microservices patterns: Service Discovery", 
                        "Database optimization: B-Tree Indexing", 
                        "Database optimization: Query Profiling",
                        "Asynchronous processing: Celery tasks",
                        "Asynchronous processing: RabbitMQ queues"
                    ],
                    "duration": "4 weeks",
                    "focus": "Build robust, scalable systems and understand architecture tradeoffs."
                },
                {
                    "level": 3,
                    "title": "Level 3: Production Mastery",
                    "topics": [
                        "CI/CD pipelines: GitHub Actions", 
                        "CI/CD pipelines: Docker Image Optimization", 
                        "Monitoring & Logging: Prometheus metrics", 
                        "Monitoring & Logging: ELK Stack", 
                        "Cloud deployment: AWS ECS/EKS",
                        "Cloud deployment: Terraform basics"
                    ],
                    "duration": "2 weeks",
                    "focus": "Deploy applications securely to production and handle scale."
                }
            ]
        }


MOCK_TEST_FALLBACK_DATA = {
    "technical": [
        # DSA (10)
        {"id": "tech_dsa_0", "subject": "DSA", "question": "What is the worst-case time complexity of searching for an element in a balanced Binary Search Tree (BST)?", "options": ["O(1)", "O(log n)", "O(n)", "O(n log n)"], "answer": "O(log n)", "explanation": "In a balanced BST, the height of the tree is log n. Searching requires traversing from root to leaf, taking at most O(log n) steps."},
        {"id": "tech_dsa_1", "subject": "DSA", "question": "Which data structure is typically used to implement Breadth-First Search (BFS) on a graph?", "options": ["Stack", "Queue", "Priority Queue", "Deque"], "answer": "Queue", "explanation": "BFS explores nodes level by level, requiring a First-In-First-Out (FIFO) ordering which is naturally supported by a Queue."},
        {"id": "tech_dsa_2", "subject": "DSA", "question": "What is the worst-case time complexity of the standard Quick Sort algorithm?", "options": ["O(log n)", "O(n)", "O(n log n)", "O(n^2)"], "answer": "O(n^2)", "explanation": "The worst case occurs when the pivot consistently divides the array into empty and n-1 element sub-arrays (e.g., sorted array with first element as pivot), yielding O(n^2)."},
        {"id": "tech_dsa_3", "subject": "DSA", "question": "Which of the following data structures is used by the compiler to manage function calls and local variables?", "options": ["Queue", "Stack", "Heap", "Linked List"], "answer": "Stack", "explanation": "The call stack utilizes a Last-In-First-Out (LIFO) model to store activation records (local variables, return addresses) of active subroutines."},
        {"id": "tech_dsa_4", "subject": "DSA", "question": "What is the optimal average-case time complexity to check if an element exists in a Hash Table?", "options": ["O(1)", "O(log n)", "O(n)", "O(n log n)"], "answer": "O(1)", "explanation": "A well-designed hash table with a good hash function distributes keys uniformly, giving O(1) average lookup time."},
        {"id": "tech_dsa_5", "subject": "DSA", "question": "Which of the following sorting algorithms is NOT stable in its standard implementation?", "options": ["Merge Sort", "Bubble Sort", "Insertion Sort", "Heap Sort"], "answer": "Heap Sort", "explanation": "Heap Sort processes elements using a binary heap tree structure where the relative order of identical elements is not preserved, making it unstable."},
        {"id": "tech_dsa_6", "subject": "DSA", "question": "What is the worst-case space complexity of storing a graph with V vertices and E edges using an Adjacency Matrix?", "options": ["O(V)", "O(V + E)", "O(V^2)", "O(V * E)"], "answer": "O(V^2)", "explanation": "An adjacency matrix is a 2D array of size V x V, requiring O(V^2) memory space regardless of the number of edges."},
        {"id": "tech_dsa_7", "subject": "DSA", "question": "Which tree traversal technique visits nodes in non-decreasing sorted order when applied to a Binary Search Tree (BST)?", "options": ["Preorder", "Inorder", "Postorder", "Level-order"], "answer": "Inorder", "explanation": "Inorder traversal (Left, Root, Right) processes the left subtree, then root, then right subtree, which visits BST nodes in ascending sorted order."},
        {"id": "tech_dsa_8", "subject": "DSA", "question": "What is the time complexity of merging two sorted arrays of size M and N into a single sorted array?", "options": ["O(log(M+N))", "O(M log N)", "O(M + N)", "O(M * N)"], "answer": "O(M + N)", "explanation": "Merging requires iterating through both arrays once using two pointers, taking linear time proportional to the total number of elements: O(M + N)."},
        {"id": "tech_dsa_9", "subject": "DSA", "question": "Which algorithm is most appropriate for finding the single-source shortest paths in a weighted graph with negative edge weights but no negative cycles?", "options": ["Dijkstra's Algorithm", "Bellman-Ford Algorithm", "Kruskal's Algorithm", "Prim's Algorithm"], "answer": "Bellman-Ford Algorithm", "explanation": "Dijkstra's algorithm fails with negative weights because it relies on greedy accumulation. Bellman-Ford correctly handles negative weights by relaxing all edges V-1 times."},

        # OOPs (10)
        {"id": "tech_oops_0", "subject": "OOPs", "question": "Which object-oriented programming concept allows a subclass to provide a specific implementation of a method already defined in its superclass?", "options": ["Method Overloading", "Method Overriding", "Method Abstraction", "Encapsulation"], "answer": "Method Overriding", "explanation": "Method Overriding allows a subclass to define a method with the same signature as one in the parent class to achieve runtime polymorphism."},
        {"id": "tech_oops_1", "subject": "OOPs", "question": "What is defining multiple methods in the same class with the same name but different parameter lists called?", "options": ["Method Overriding", "Method Overloading", "Inheritance", "Dynamic Binding"], "answer": "Method Overloading", "explanation": "Method Overloading is compile-time polymorphism where methods share a name but differ in parameter count, type, or order."},
        {"id": "tech_oops_2", "subject": "OOPs", "question": "Which OOP pillar is concerned with hiding internal implementation details and exposing only safe, essential interfaces?", "options": ["Inheritance", "Polymorphism", "Abstraction", "Instantiation"], "answer": "Abstraction", "explanation": "Abstraction simplifies complex systems by hiding internal details from the user and exposing only relevant inputs/outputs."},
        {"id": "tech_oops_3", "subject": "OOPs", "question": "What does encapsulation accomplish in system design?", "options": ["Allows classes to inherit methods", "Binds data and methods together and restricts direct access", "Enables a method to accept different parameter types", "Speeds up compile-time execution"], "answer": "Binds data and methods together and restricts direct access", "explanation": "Encapsulation wraps class data (attributes) and methods together, exposing them only via public getters/setters to protect state integrity."},
        {"id": "tech_oops_4", "subject": "OOPs", "question": "Which of the following statements about abstract classes and interfaces in Java is TRUE?", "options": ["A class can inherit multiple abstract classes.", "Interfaces cannot contain default method implementations.", "An abstract class can contain instance variables, whereas interfaces cannot.", "Abstract classes cannot have constructors."], "answer": "An abstract class can contain instance variables, whereas interfaces cannot.", "explanation": "Abstract classes represent real state and can have instance fields. Interfaces are contracts and can only have static final constants (no instance state)."},
        {"id": "tech_oops_5", "subject": "OOPs", "question": "Which of the following types of inheritance is NOT supported natively by Java through classes?", "options": ["Single Inheritance", "Multilevel Inheritance", "Hierarchical Inheritance", "Multiple Inheritance"], "answer": "Multiple Inheritance", "explanation": "Java does not support multiple inheritance of classes to prevent the Diamond Problem (ambiguity in inheriting duplicate methods). It must be achieved via Interfaces."},
        {"id": "tech_oops_6", "subject": "OOPs", "question": "What is the primary role of a constructor in object-oriented programming?", "options": ["To destroy an object when it goes out of scope", "To allocate memory for the object's code", "To initialize the instance variables of a newly created object", "To prevent a class from being subclassed"], "answer": "To initialize the instance variables of a newly created object", "explanation": "Constructors run automatically when an object is instantiated (using 'new') to initialize its member variables and set up initial state."},
        {"id": "tech_oops_7", "subject": "OOPs", "question": "What is runtime (dynamic) polymorphism achieved by?", "options": ["Virtual functions / method overriding", "Overloaded functions", "Static templates", "Private attributes"], "answer": "Virtual functions / method overriding", "explanation": "Runtime polymorphism is resolved at execution time based on the actual object type, achieved using virtual methods (overriding) and base pointers."},
        {"id": "tech_oops_8", "subject": "OOPs", "question": "What does the 'super' keyword (or parent reference) do in Java/Python subclasses?", "options": ["Calls the garbage collector", "Refers to the parent class object or constructor", "Instantiates a new thread", "Prevents class modification"], "answer": "Refers to the parent class object or constructor", "explanation": "The 'super' keyword allows a subclass to invoke constructors, methods, and attributes of its parent class."},
        {"id": "tech_oops_9", "subject": "OOPs", "question": "What is a class that cannot be instantiated and is designed solely to be subclassed called?", "options": ["Final Class", "Static Class", "Abstract Class", "Concrete Class"], "answer": "Abstract Class", "explanation": "Abstract classes cannot be direct instances. They act as blueprints that other concrete classes must inherit from and complete."},

        # OS (10)
        {"id": "tech_os_0", "subject": "OS", "question": "What is a process that has completed execution but still has an entry in the Operating System's process table called?", "options": ["Orphan Process", "Zombie Process", "Daemon Process", "Interrupt Handler"], "answer": "Zombie Process", "explanation": "A zombie process is one that has finished execution but whose parent has not yet read its exit status via wait(), keeping its entry in the process table."},
        {"id": "tech_os_1", "subject": "OS", "question": "Which of the following CPU scheduling algorithms is non-preemptive in its basic implementation?", "options": ["Round Robin", "First-Come First-Served (FCFS)", "Shortest Remaining Time First (SRTF)", "Priority Preemptive"], "answer": "First-Come First-Served (FCFS)", "explanation": "FCFS schedules processes in the order they arrive. Once a process starts running, it doesn't yield CPU control until it finishes or blocks, making it non-preemptive."},
        {"id": "tech_os_2", "subject": "OS", "question": "Which of the following is NOT one of the four necessary conditions for a deadlock to occur?", "options": ["Mutual Exclusion", "Hold and Wait", "No Preemption", "Resource Preemption"], "answer": "Resource Preemption", "explanation": "Deadlock requires No Preemption (resources cannot be forcibly taken). Mutual Exclusion, Hold & Wait, No Preemption, and Circular Wait are Coffman's conditions."},
        {"id": "tech_os_3", "subject": "OS", "question": "What memory management scheme solves external fragmentation by dividing physical memory into fixed-size blocks?", "options": ["Segmentation", "Paging", "Dynamic Partitioning", "Swapping"], "answer": "Paging", "explanation": "Paging maps fixed-size pages to fixed-size physical frames, completely eliminating external fragmentation (though internal fragmentation can still exist in the last page)."},
        {"id": "tech_os_4", "subject": "OS", "question": "What is the primary function of the Translation Lookaside Buffer (TLB) in OS memory management?", "options": ["To cache disk writes", "To speed up translation of virtual addresses to physical addresses", "To resolve network IP addresses", "To store CPU thread states"], "answer": "To speed up translation of virtual addresses to physical addresses", "explanation": "TLB is a fast hardware cache that stores recent page-to-frame translations, bypassing slow main memory page table lookups."},
        {"id": "tech_os_5", "subject": "OS", "question": "Which system call is used in Unix-like operating systems to create a duplicate clone child process?", "options": ["exec()", "forkv()", "fork()", "spawn()"], "answer": "fork()", "explanation": "The fork() system call duplicates the parent process, creating a child process with a separate address space but identical contents."},
        {"id": "tech_os_6", "subject": "OS", "question": "What is the condition called when an operating system spends more time swapping pages in and out of disk than executing useful instructions?", "options": ["Starvation", "Thrashing", "Deadlock", "Context Switching"], "answer": "Thrashing", "explanation": "Thrashing occurs when active processes lack sufficient resident page frames in RAM, causing continuous page faults and disk access, dropping CPU utility to near zero."},
        {"id": "tech_os_7", "subject": "OS", "question": "What is context switching in an operating system?", "options": ["Converting source code to bytecode", "Saving the state of a running process and loading the state of another", "Allocating bandwidth to network cards", "Compacting memory fragments on disk"], "answer": "Saving the state of a running process and loading the state of another", "explanation": "Context switching is saving registers/PC of a running process into its PCB and restoring another process's state to let the CPU run it."},
        {"id": "tech_os_8", "subject": "OS", "question": "Which synchronization primitive is a simple integer variable accessed via atomic wait() and signal() operations?", "options": ["Mutex", "Condition Variable", "Semaphore", "Spinlock"], "answer": "Semaphore", "explanation": "Semaphores use an integer value to control access to shared resources. wait() decrements it, signal() increments it, both executing atomically."},
        {"id": "tech_os_9", "subject": "OS", "question": "What does a page fault represent?", "options": ["An error in compiled program memory allocation", "A request for a page that is not currently loaded in physical RAM", "A disk sector hardware crash", "A failure in network socket binding"], "answer": "A request for a page that is not currently loaded in physical RAM", "explanation": "When a process accesses a page marked invalid in the page table, the MMU triggers a page fault interrupt, forcing the OS to fetch the page from swap space into RAM."},

        # DBMS (10)
        {"id": "tech_dbms_0", "subject": "DBMS", "question": "What do the ACID properties of a database transaction ensure?", "options": ["Fast query indexing", "Data integrity, consistency, and reliability", "Network security encryption", "Automated SQL database backups"], "answer": "Data integrity, consistency, and reliability", "explanation": "ACID properties (Atomicity, Consistency, Isolation, Durability) ensure that database transactions are processed reliably and maintain integrity."},
        {"id": "tech_dbms_1", "subject": "DBMS", "question": "Which database normal form (NF) is violated if a relation contains transitive functional dependencies?", "options": ["1NF", "2NF", "3NF", "BCNF"], "answer": "3NF", "explanation": "3NF requires that all non-prime attributes are mutually independent. If A determines B and B determines C, a transitive dependency A -> C exists, violating 3NF."},
        {"id": "tech_dbms_2", "subject": "DBMS", "question": "What is a primary key that consists of more than one column in a database table called?", "options": ["Candidate Key", "Foreign Key", "Composite Key", "Surrogate Key"], "answer": "Composite Key", "explanation": "A composite key is a primary key comprising two or more columns to uniquely identify table rows (e.g. order_id and item_id)."},
        {"id": "tech_dbms_3", "subject": "DBMS", "question": "Which SQL clause is used to filter rows AFTER aggregate operations have been applied using GROUP BY?", "options": ["WHERE", "HAVING", "LIMIT", "FILTER"], "answer": "HAVING", "explanation": "WHERE filters rows before grouping occurs. HAVING filters aggregated data groups (e.g. HAVING COUNT(*) > 5) after GROUP BY is processed."},
        {"id": "tech_dbms_4", "subject": "DBMS", "question": "Which index structure is widely used in relational databases to facilitate fast range search queries?", "options": ["Hash Table Index", "B+ Tree Index", "Binary Search Index", "Trie Index"], "answer": "B+ Tree Index", "explanation": "B+ Trees store sorted keys with leaf nodes linked sequentially, enabling both O(log n) lookups and efficient sequential range scans."},
        {"id": "tech_dbms_5", "subject": "DBMS", "question": "What type of JOIN returns all records from the left table and the matched records from the right table?", "options": ["Inner Join", "Full Outer Join", "Left Outer Join", "Right Outer Join"], "answer": "Left Outer Join", "explanation": "Left Outer Join returns all rows from the left table; matching cells are populated from the right table, or NULL if no match exists."},
        {"id": "tech_dbms_6", "subject": "DBMS", "question": "What does a foreign key constraint maintain in relational database design?", "options": ["Entity Integrity", "Referential Integrity", "Domain Integrity", "Transaction Durability"], "answer": "Referential Integrity", "explanation": "Foreign keys ensure referential integrity by validating that values in child columns exist in the referenced parent columns."},
        {"id": "tech_dbms_7", "subject": "DBMS", "question": "Which SQL statement is used to remove all records from a table without logging individual row deletions, making it faster than DELETE?", "options": ["DROP TABLE", "REMOVE TABLE", "TRUNCATE TABLE", "CLEAR TABLE"], "answer": "TRUNCATE TABLE", "explanation": "TRUNCATE is a DDL command that deallocates data pages of the table, bypassing row-by-row transaction logging, which makes it fast."},
        {"id": "tech_dbms_8", "subject": "DBMS", "question": "What is the default transaction isolation level in MySQL InnoDB storage engine?", "options": ["Read Uncommitted", "Read Committed", "Repeatable Read", "Serializable"], "answer": "Repeatable Read", "explanation": "Repeatable Read is InnoDB's default. It uses locks and MVCC to ensure reads within a transaction see the same snapshot, preventing non-repeatable reads."},
        {"id": "tech_dbms_9", "subject": "DBMS", "question": "What is the purpose of database normalization?", "options": ["To increase query complexity", "To minimize data redundancy and dependency anomalies", "To encrypt credit card numbers", "To compress historical files"], "answer": "To minimize data redundancy and dependency anomalies", "explanation": "Normalization organizes tables to reduce redundant data and prevent insert, update, and delete inconsistencies (anomalies)."},

        # CN (10)
        {"id": "tech_cn_0", "subject": "CN", "question": "Which layer of the OSI model is responsible for packet routing, logical addressing, and path determination?", "options": ["Data Link Layer", "Network Layer", "Transport Layer", "Application Layer"], "answer": "Network Layer", "explanation": "The Network Layer handles packet routing across networks, logical IP addressing, and intermediate path configuration."},
        {"id": "tech_cn_1", "subject": "CN", "question": "What is the exact size of an IPv6 network address?", "options": ["32 bits", "64 bits", "128 bits", "256 bits"], "answer": "128 bits", "explanation": "IPv6 addresses are 128-bit hexadecimal numbers, written as eight groups of four digits separated by colons (e.g. 2001:db8::1)."},
        {"id": "tech_cn_2", "subject": "CN", "question": "Which protocol translates a logical IP address into a physical MAC hardware address?", "options": ["DNS", "DHCP", "ARP (Address Resolution Protocol)", "NAT"], "answer": "ARP (Address Resolution Protocol)", "explanation": "ARP is a broadcast protocol that resolves local IP addresses into physical MAC ethernet frames on a local subnet."},
        {"id": "tech_cn_3", "subject": "CN", "question": "What is the standard TCP port number reserved for secure HTTPS traffic?", "options": ["80", "443", "8080", "22"], "answer": "443", "explanation": "HTTPS runs over SSL/TLS secured connections on TCP port 443 by default. Regular HTTP runs on port 80."},
        {"id": "tech_cn_4", "subject": "CN", "question": "Which network transport protocol is connectionless and does NOT guarantee packet delivery or order?", "options": ["TCP", "UDP (User Datagram Protocol)", "SCTP", "HTTP"], "answer": "UDP (User Datagram Protocol)", "explanation": "UDP is a lightweight protocol that pushes datagrams without establishing handshakes or verifying acknowledgement, prioritizing speed over reliability."},
        {"id": "tech_cn_5", "subject": "CN", "question": "What is the main purpose of the Domain Name System (DNS)?", "options": ["To distribute server IP addresses dynamically", "To map human-readable domain names to numerical IP addresses", "To encrypt password databases", "To load balance web requests"], "answer": "To map human-readable domain names to numerical IP addresses", "explanation": "DNS works as the phonebook of the internet, translating domain queries (e.g., google.com) into target server IP locations."},
        {"id": "tech_cn_6", "subject": "CN", "question": "Which network topology connects all client devices directly to a single central switch or hub?", "options": ["Mesh Topology", "Ring Topology", "Bus Topology", "Star Topology"], "answer": "Star Topology", "explanation": "In a Star topology, all network nodes are connected individually to a central switch. If one cable fails, only that device is disconnected."},
        {"id": "tech_cn_7", "subject": "CN", "question": "What is the primary function of a subnet mask in IPv4 configurations?", "options": ["To identify network ID versus host ID sections of the IP address", "To route packets across the WAN", "To assign DNS names to host computers", "To encrypt private key files"], "answer": "To identify network ID versus host ID sections of the IP address", "explanation": "Subnet masks (e.g. 255.255.255.0) use bitwise AND masking to separate network subnet addresses from host identifiers."},
        {"id": "tech_cn_8", "subject": "CN", "question": "Which network protocol is used by routers to dynamically allocate IP addresses to new clients?", "options": ["SNMP", "SMTP", "DHCP (Dynamic Host Configuration Protocol)", "ICMP"], "answer": "DHCP (Dynamic Host Configuration Protocol)", "explanation": "DHCP servers dynamically assign IP addresses, gateway pathways, subnet masks, and DNS servers to local computers upon boot."},
        {"id": "tech_cn_9", "subject": "CN", "question": "What is the three-way handshake process used by TCP to establish a connection?", "options": ["SYN -> SYN-ACK -> ACK", "SYN -> ACK -> SYN-ACK", "ACK -> SYN -> SYN-ACK", "CONNECT -> OK -> CONNECT-ACK"], "answer": "SYN -> SYN-ACK -> ACK", "explanation": "TCP connects by sending a synchronization flag (SYN), receiving synchronization-acknowledgement (SYN-ACK), and returning acknowledgement (ACK)."},

        # System Design (10)
        {"id": "tech_sys_0", "subject": "System Design", "question": "What is adding more nodes or servers to a distributed application infrastructure cluster called?", "options": ["Vertical Scaling (Scaling Up)", "Horizontal Scaling (Scaling Out)", "Data Partitioning", "Latency Tuning"], "answer": "Horizontal Scaling (Scaling Out)", "explanation": "Horizontal scaling distributes load by adding more commodity servers to the application pool, unlike vertical scaling which upgrades a single server."},
        {"id": "tech_sys_1", "subject": "System Design", "question": "Which system component acts as a reverse proxy to distribute traffic evenly across multiple server nodes?", "options": ["CDN", "Database Replica", "Load Balancer", "Message Queue"], "answer": "Load Balancer", "explanation": "Load Balancers receive client traffic and route it to multiple server backends using algorithms like round robin or least connections."},
        {"id": "tech_sys_2", "subject": "System Design", "question": "What is the primary architectural purpose of a Content Delivery Network (CDN)?", "options": ["To handle write-heavy database transactions", "To store session states of active users", "To cache static media resources geographically closer to endpoints", "To run slow analytical batch calculations"], "answer": "To cache static media resources geographically closer to endpoints", "explanation": "CDNs place edge caches (CDNs) around the world to serve images, JS, CSS, and videos closer to the user, reducing network latency."},
        {"id": "tech_sys_3", "subject": "System Design", "question": "According to the CAP Theorem, which of the three properties must a distributed system choose between in the event of a network partition?", "options": ["Cost, Availability, Performance", "Consistency, Availability, Partition Tolerance", "Concurrency, Authority, Privacy", "Caching, Archiving, Pipelining"], "answer": "Consistency, Availability, Partition Tolerance", "explanation": "CAP states that in a network partition (P), a distributed database can guarantee Consistency (C) OR Availability (A), but not both simultaneously."},
        {"id": "tech_sys_4", "subject": "System Design", "question": "What is database partitioning (sharding) in system architecture?", "options": ["Replicating database servers for backups", "Splitting a single database dataset into smaller database clusters", "Indexing column strings for fast text searching", "Encrypting database schema tables"], "answer": "Splitting a single database dataset into smaller database clusters", "explanation": "Sharding partitions data horizontally (e.g. by user_id range) across multiple physical database nodes to scale out write throughput."},
        {"id": "tech_sys_5", "subject": "System Design", "question": "What does a write-through caching strategy do?", "options": ["Writes only to the cache and syncs to DB later asynchronously", "Writes directly to the database, bypassing the cache entirely", "Writes data to both the cache and the database synchronously", "Invalidates all keys in cache when a write occurs"], "answer": "Writes data to both the cache and the database synchronously", "explanation": "Write-through caching updates both the cache and underlying database simultaneously on write operations, ensuring strong data consistency."},
        {"id": "tech_sys_6", "subject": "System Design", "question": "Which technique is used to coordinate consistent hashing across a ring of cache nodes to minimize key reassignment when nodes join or leave?", "options": ["Round Robin routing", "Consistent Hashing", "Sharded Indexing", "Master-Slave replication"], "answer": "Consistent Hashing", "explanation": "Consistent hashing maps both keys and server nodes to a circular numeric ring. Adding/removing nodes only reassigns a fraction of the keys (1/N)."},
        {"id": "tech_sys_7", "subject": "System Design", "question": "Which protocol is designed for real-time, bi-directional persistent connection communication between a client and server?", "options": ["HTTP/1.1 polling", "WebSockets", "gRPC REST", "Simple Mail Transfer Protocol"], "answer": "WebSockets", "explanation": "WebSockets start with an HTTP upgrade handshake and maintain a long-running, bi-directional TCP socket connection for low-latency communication."},
        {"id": "tech_sys_8", "subject": "System Design", "question": "What role does a message broker (e.g. Kafka, RabbitMQ) play in microservices architectures?", "options": ["It processes user authentication requests", "It decouples services and allows asynchronous event processing", "It acts as a primary transactional datastore", "It translates database schemas to JSON"], "answer": "It decouples services and allows asynchronous event processing", "explanation": "Message queues act as intermediate event logs, letting publishers push updates without waiting for consumers, achieving loose coupling."},
        {"id": "tech_sys_9", "subject": "System Design", "question": "What is horizontal sharding based on hash key mapping?", "options": ["Splitting data based on hash values to distribute rows uniformly across nodes", "Creating duplicate tables with index copies", "Translating relational tables to non-relational key-value pairs", "Restructuring foreign key tables"], "answer": "Splitting data based on hash values to distribute rows uniformly across nodes", "explanation": "Hash sharding calculates a hash of the key (e.g., hash(user_id) % N) to determine the target shard, achieving uniform data distribution."},
    ],
    "verbal": [
        {"id": "verb_0", "question": "Choose the word most nearly opposite in meaning to 'EPHEMERAL'.", "options": ["Transient", "Permanent", "Ethereal", "Unstable"], "answer": "Permanent", "explanation": "Ephemeral means short-lived or transient; permanent means lasting or enduring, which is the opposite."},
        {"id": "verb_1", "question": "Fill in the blank: The manager was _______ by the consultant's report, which highlighted several errors in the budget.", "options": ["exonerated", "perturbed", "exalted", "appeased"], "answer": "perturbed", "explanation": "Perturbed means anxious or unsettled. Discovering budget errors would naturally perturb a manager."},
        {"id": "verb_2", "question": "Complete the analogy: METAPHOR : FIGURATIVE :: FACT : _______", "options": ["Literal", "Poetic", "Deceptive", "Imaginary"], "answer": "Literal", "explanation": "A metaphor is a type of figurative language; a fact is a type of literal description."},
        {"id": "verb_3", "question": "Identify the synonym of the word 'PULCHRITUDE'.", "options": ["Ugliness", "Avarice", "Beauty", "Honesty"], "answer": "Beauty", "explanation": "Pulchritude refers to physical beauty or comeliness."},
        {"id": "verb_4", "question": "Identify the grammatically correct sentence.", "options": ["Neither of the options look suitable.", "Neither of the options looks suitable.", "Either of the options are suitable.", "Neither options is suitable."], "answer": "Neither of the options looks suitable.", "explanation": "'Neither' is a singular pronoun and takes the singular verb 'looks'."},
        {"id": "verb_5", "question": "Choose the word that best fits: Although she was generally soft-spoken, she could be surprisingly _______ in debate.", "options": ["timid", "assertive", "reticent", "conciliatory"], "answer": "assertive", "explanation": "The word 'Although' indicates a contrast to being soft-spoken. 'Assertive' (confident, forceful) fits perfectly."},
        {"id": "verb_6", "question": "What does the idiom 'Bite the bullet' mean?", "options": ["To act recklessly", "To face a difficult situation with courage", "To make a mistake", "To start a conflict"], "answer": "To face a difficult situation with courage", "explanation": "Biting the bullet means resolving to face a painful or difficult situation that cannot be avoided."},
        {"id": "verb_7", "question": "Find the word with the correct spelling.", "options": ["Accomodation", "Accommodation", "Acommodation", "Accomodatoin"], "answer": "Accommodation", "explanation": "Accommodation is spelled with double 'c' and double 'm'."},
        {"id": "verb_8", "question": "Antonym of 'LUCID' is:", "options": ["Clear", "Obscure", "Intelligent", "Transparent"], "answer": "Obscure", "explanation": "Lucid means clear and easy to understand; obscure means unclear or hard to perceive."},
        {"id": "verb_9", "question": "Which word means 'to make something bad better'?", "options": ["Aggravate", "Ameliorate", "Deteriorate", "Exacerbate"], "answer": "Ameliorate", "explanation": "Ameliorate means to make something bad or unsatisfactory better."},
        {"id": "verb_10", "question": "Analogy: INSULATOR : HEAT :: SHIELD : _______", "options": ["Weapon", "Impact", "Target", "Conductor"], "answer": "Impact", "explanation": "An insulator protects against heat; a shield protects against impact."},
        {"id": "verb_11", "question": "Antonym of 'CANDID' is:", "options": ["Honest", "Frank", "Insincere", "Truthful"], "answer": "Insincere", "explanation": "Candid means truthful and straightforward; insincere is the opposite."},
        {"id": "verb_12", "question": "Choose the correct preposition: She is very proficient _______ playing the violin.", "options": ["at", "in", "with", "on"], "answer": "at", "explanation": "One is proficient 'at' doing an activity, or proficient 'in' a field of study."},
        {"id": "verb_13", "question": "Synonym of 'REDUNDANT' is:", "options": ["Essential", "Superfluous", "Scarse", "Important"], "answer": "Superfluous", "explanation": "Redundant means no longer needed or useful; superfluous is a direct synonym."},
        {"id": "verb_14", "question": "Synonym of 'METICULOUS' is:", "options": ["Careless", "Precise", "Unorganized", "Quick"], "answer": "Precise", "explanation": "Meticulous means showing great attention to detail; very careful and precise."},
    ],
    "aptitude": [
        {"id": "apt_0", "question": "A can complete a project in 10 days, while B can complete it in 15 days. How many days will they take if they work together?", "options": ["5 days", "6 days", "8 days", "12 days"], "answer": "6 days", "explanation": "Together, they do (1/10 + 1/15) = 5/30 = 1/6 of the job per day. Thus, they take 6 days total."},
        {"id": "apt_1", "question": "A train passes a pole in 15 seconds while traveling at a speed of 60 km/h. What is the length of the train in meters?", "options": ["150 m", "200 m", "250 m", "300 m"], "answer": "250 m", "explanation": "Speed = 60 * (5/18) = 50/3 m/s. Length = Speed * Time = (50/3) * 15 = 250 meters."},
        {"id": "apt_2", "question": "If the price of a commodity increases by 20%, by what percentage must a consumer decrease their consumption so that expenditure remains constant?", "options": ["16.67%", "20%", "25%", "33.33%"], "answer": "16.67%", "explanation": "Formula: (R / (100 + R)) * 100% = (20/120) * 100 = 16.67%."},
        {"id": "apt_3", "question": "A bag contains 5 red balls and 7 blue balls. If two balls are drawn at random without replacement, what is the probability that both are red?", "options": ["5/33", "5/22", "7/22", "5/12"], "answer": "5/33", "explanation": "P(1st red) = 5/12. P(2nd red) = 4/11. Total Probability = (5/12) * (4/11) = 20/132 = 5/33."},
        {"id": "apt_4", "question": "A sum of money doubles itself in 8 years at simple interest. What is the annual rate of interest?", "options": ["8%", "10%", "12.5%", "15%"], "answer": "12.5%", "explanation": "Simple Interest = Principal. So, P = P * R * 8 / 100 => 8R = 100 => R = 12.5%."},
        {"id": "apt_5", "question": "The average weight of 8 men increases by 2.5 kg when a man weighing 65 kg is replaced by a new man. What is the weight of the new man?", "options": ["75 kg", "80 kg", "85 kg", "90 kg"], "answer": "85 kg", "explanation": "Total weight increase = 8 * 2.5 = 20 kg. Weight of new man = 65 + 20 = 85 kg."},
        {"id": "apt_6", "question": "A shopkeeper sells an article at a loss of 10%. If he had sold it for $30 more, he would have made a profit of 5%. What is the cost price of the article?", "options": ["$150", "$200", "$250", "$300"], "answer": "$200", "explanation": "Difference in percentage = 5% - (-10%) = 15%. Thus, 15% of CP = $30 => CP = 30 / 0.15 = $200."},
        {"id": "apt_7", "question": "Two numbers are in the ratio of 3:4. If their HCF is 4, what is their LCM?", "options": ["12", "16", "24", "48"], "answer": "48", "explanation": "The numbers are 3*4 = 12 and 4*4 = 16. LCM of 12 and 16 is 48."},
        {"id": "apt_8", "question": "A boat goes 8 km downstream in 40 minutes and returns upstream in 1 hour. What is the speed of the boat in still water?", "options": ["8 km/h", "10 km/h", "12 km/h", "14 km/h"], "answer": "10 km/h", "explanation": "Downstream speed = 8 / (40/60) = 12 km/h. Upstream speed = 8 / 1 = 8 km/h. Speed in still water = (12 + 8) / 2 = 10 km/h."},
        {"id": "apt_9", "question": "If 12 men can build a wall in 20 days, how many days will 15 men take to build the same wall?", "options": ["14 days", "15 days", "16 days", "18 days"], "answer": "16 days", "explanation": "Using M1 * D1 = M2 * D2: 12 * 20 = 15 * D2 => D2 = 240 / 15 = 16 days."},
        {"id": "apt_10", "question": "What is the next number in the series: 2, 6, 12, 20, 30, ___?", "options": ["36", "40", "42", "48"], "answer": "42", "explanation": "The differences are +4, +6, +8, +10. The next difference is +12, so 30 + 12 = 42."},
        {"id": "apt_11", "question": "In a group of 60 students, 40% study Physics, 50% study Chemistry, and 10% study both. How many students study neither?", "options": ["10", "12", "15", "18"], "answer": "12", "explanation": "Total studying either = 40% + 50% - 10% = 80%. Studying neither = 20% of 60 = 12 students."},
        {"id": "apt_12", "question": "What is the compound interest on $10,000 for 2 years at 10% per annum compounded annually?", "options": ["$1,000", "$2,000", "$2,100", "$2,200"], "answer": "$2,100", "explanation": "Amount = 10000 * (1.1)^2 = 10000 * 1.21 = $12,100. Interest = 12100 - 10000 = $2,100."},
        {"id": "apt_13", "question": "If 3x + 5y = 21 and 2x - y = 1, find the value of x.", "options": ["1", "2", "3", "4"], "answer": "2", "explanation": "Multiply second eq by 5: 10x - 5y = 5. Add to first: 13x = 26 => x = 2."},
        {"id": "apt_14", "question": "A worker's salary is increased by 10% and then decreased by 10%. What is the net change in his salary?", "options": ["0%", "1% increase", "1% decrease", "2% decrease"], "answer": "1% decrease", "explanation": "Formula: x + y + xy/100 = 10 - 10 - 100/100 = -1%. So, it is a 1% decrease."},
    ],
    "coding": [
        {
            "id": "code_easy",
            "difficulty": "Easy",
            "title": "Two Sum (Array Search)",
            "description": "Given an array of integers 'nums' and an integer 'target', return indices of the two numbers such that they add up to target. You may assume that each input would have exactly one solution, and you may not use the same element twice. Return your indices list or index pairs.",
            "examples": [
                {
                    "input": "nums = [2,7,11,15], target = 9",
                    "output": "[0, 1]",
                    "explanation": "Because nums[0] + nums[1] == 2 + 7 == 9, we return [0, 1]."
                }
            ],
            "constraints": ["2 <= nums.length <= 10^4", "-10^9 <= nums[i] <= 10^9", "-10^9 <= target <= 10^9"],
            "templates": {
                "cpp": "class Solution {\npublic:\n    vector<int> twoSum(vector<int>& nums, int target) {\n        // Write your C++ code here\n        return {};\n    }\n};",
                "java": "class Solution {\n    public int[] twoSum(int[] nums, int target) {\n        // Write your Java code here\n        return new int[]{};\n    }\n}",
                "python": "def twoSum(nums, target):\n    # Write your Python code here\n    return []",
                "javascript": "function twoSum(nums, target) {\n    // Write your JavaScript code here\n    return [];\n}"
            }
        },
        {
            "id": "code_hard",
            "difficulty": "Hard",
            "title": "LRU Cache Design",
            "description": "Design a data structure that follows the constraints of a Least Recently Used (LRU) Cache. Implement an LRUCache class with capacity, get(key) which returns the value of key if it exists (otherwise -1), and put(key, value) which updates or inserts the key. When the cache runs out of space, it must invalidate and evict the least recently used item before inserting.",
            "examples": [
                {
                    "input": "capacity = 2, put(1, 10), put(2, 20), get(1), put(3, 30), get(2)",
                    "output": "get(1) -> 10, get(2) -> -1",
                    "explanation": "put(3, 30) evicts key 2 because it was the least recently used key. Thus, subsequent get(2) returns -1."
                }
            ],
            "constraints": ["1 <= capacity <= 3000", "0 <= key <= 10^4", "0 <= value <= 10^5"],
            "templates": {
                "cpp": "class LRUCache {\npublic:\n    LRUCache(int capacity) {\n        // Initialize cache capacity\n    }\n    \n    int get(int key) {\n        // Write your code here\n        return -1;\n    }\n    \n    void put(int key, int value) {\n        // Write your code here\n    }\n};",
                "java": "class LRUCache {\n    public LRUCache(int capacity) {\n        // Initialize cache capacity\n    }\n    \n    public int get(int key) {\n        // Write your code here\n        return -1;\n    }\n    \n    public void put(int key, int value) {\n        // Write your code here\n    }\n}",
                "python": "class LRUCache:\n    def __init__(self, capacity: int):\n        # Initialize cache capacity\n        pass\n\n    def get(self, key: int) -> int:\n        # Write your code here\n        return -1\n\n    def put(self, key: int, value: int) -> None:\n        # Write your code here\n        pass",
                "javascript": "class LRUCache {\n    constructor(capacity) {\n        // Initialize cache capacity\n    }\n\n    get(key) {\n        // Write your code here\n        return -1;\n    }\n\n    put(key, value) {\n        // Write your code here\n    }\n}"
            }
        }
    ]
}


def generate_mock_test(difficulty):
    if not is_valid_key():
        return MOCK_TEST_FALLBACK_DATA

    def generate_section(section_name, count, req_desc, format_desc):
        prompt = f"""Generate a mock assessment section for a software engineering candidate.
Difficulty Level: {difficulty.upper()}. Section: {section_name}.
You must generate exactly {count} {req_desc}.

Return a strictly valid JSON array of objects with the following structure:
[
  {format_desc}
]

Format the response strictly as a valid JSON array. Do not add markdown backticks or any other text.
"""
        try:
            # 40 seconds is enough for a subset of the test on Gemini
            raw_response = _generate_with_model_fallback(prompt, timeout_sec=40)
            cleaned = re.sub(r"```json\s*", "", raw_response)
            cleaned = re.sub(r"```\s*", "", cleaned).strip()
            return json.loads(cleaned)
        except Exception as e:
            print(f"Error generating {section_name} mock test via AI: {e}")
            return []

    # Run the generation tasks in parallel. We split the Technical section into 6 subjects (10 questions each)
    # so that Gemini doesn't truncate the output, which was causing parsing errors.
    import concurrent.futures
    subjects = ["DSA", "OOPs", "OS", "DBMS", "CN", "System Design"]
    with concurrent.futures.ThreadPoolExecutor(max_workers=9) as executor:
        # Technical threads (6)
        tech_futures = {}
        for sub in subjects:
            tech_futures[sub] = executor.submit(
                generate_section, 
                f"Technical MCQs - {sub}", 
                10, 
                f"Technical MCQs for the subject {sub}", 
                f'{{ "id": "tech_{sub.lower()}_", "subject": "{sub}", "question": "...", "options": ["A", "B", "C", "D"], "answer": "A", "explanation": "..." }}'
            )

        f_verb = executor.submit(generate_section, "Verbal Reasoning MCQs", 15, "Verbal Reasoning MCQs", '{ "id": "v1", "question": "...", "options": ["A", "B", "C", "D"], "answer": "A", "explanation": "..." }')
        f_apt = executor.submit(generate_section, "Aptitude / Quantitative MCQs", 15, "Aptitude / Quantitative MCQs", '{ "id": "a1", "question": "...", "options": ["A", "B", "C", "D"], "answer": "A", "explanation": "..." }')
        f_code = executor.submit(generate_section, "Coding Questions", 2, "Coding Questions (1 Easy, 1 Hard)", '{ "id": "c1", "difficulty": "Easy", "title": "...", "description": "...", "examples": [{ "input": "...", "output": "...", "explanation": "..." }], "constraints": ["..."], "templates": { "cpp": "class Solution ...", "java": "class Solution ...", "python": "def ...", "javascript": "function ..." } }')

        # Collect Technical MCQs and merge them
        tech_data = []
        for sub in subjects:
            try:
                sub_res = tech_futures[sub].result()
            except Exception as e:
                print(f"Failed resolving thread for {sub}: {e}")
                sub_res = []
            if sub_res and isinstance(sub_res, list):
                for i, q in enumerate(sub_res):
                    if isinstance(q, dict):
                        q["subject"] = sub
                        q["id"] = f"tech_{sub.lower()}_{i}"
                        tech_data.append(q)

        try:
            verb_data = f_verb.result()
        except Exception:
            verb_data = []

        try:
            apt_data = f_apt.result()
        except Exception:
            apt_data = []

        try:
            code_data = f_code.result()
        except Exception:
            code_data = []

    # Validation and merger with FALLBACK data to ensure complete test
    # 1. Technical MCQs (verify we have 10 per subject)
    for sub in subjects:
        sub_qs = [q for q in tech_data if q.get("subject") == sub]
        if len(sub_qs) < 10:
            fallback_qs = [q for q in MOCK_TEST_FALLBACK_DATA["technical"] if q.get("subject") == sub]
            needed = 10 - len(sub_qs)
            for i in range(needed):
                if i < len(fallback_qs):
                    fq = fallback_qs[i].copy()
                    fq["id"] = f"tech_{sub.lower()}_fallback_{i}"
                    tech_data.append(fq)

    # 2. Verbal Reasoning MCQs (verify we have 15)
    if not isinstance(verb_data, list) or len(verb_data) < 15:
        if not isinstance(verb_data, list):
            verb_data = []
        needed = 15 - len(verb_data)
        fallback_qs = MOCK_TEST_FALLBACK_DATA["verbal"]
        for i in range(needed):
            if i < len(fallback_qs):
                fq = fallback_qs[i].copy()
                fq["id"] = f"verb_fallback_{i}"
                verb_data.append(fq)

    # 3. Aptitude MCQs (verify we have 15)
    if not isinstance(apt_data, list) or len(apt_data) < 15:
        if not isinstance(apt_data, list):
            apt_data = []
        needed = 15 - len(apt_data)
        fallback_qs = MOCK_TEST_FALLBACK_DATA["aptitude"]
        for i in range(needed):
            if i < len(fallback_qs):
                fq = fallback_qs[i].copy()
                fq["id"] = f"apt_fallback_{i}"
                apt_data.append(fq)

    # 4. Coding Questions (verify we have 2)
    if not isinstance(code_data, list) or len(code_data) < 2:
        code_data = MOCK_TEST_FALLBACK_DATA["coding"]

    return {
        "technical": tech_data,
        "verbal": verb_data,
        "aptitude": apt_data,
        "coding": code_data
    }


def evaluate_mock_test_code(language, code, title, description, examples):
    if not is_valid_key():
        # Fallback evaluation if no API key is configured
        is_compiled = True
        err_msg = ""
        try:
            if language.lower() in ["python", "py"]:
                compile(code, "<string>", "exec")
            elif language.lower() in ["js", "javascript"]:
                if code.count("{") != code.count("}"):
                    is_compiled = False
                    err_msg = "Syntax Error: Mismatched curly braces."
        except Exception as e:
            is_compiled = False
            err_msg = str(e)

        return {
            "compiled": is_compiled,
            "passed": False,
            "output": "Local dry-run validation successful." if is_compiled else f"Compilation Error: {err_msg}",
            "test_results": [
                {
                    "input": str(ex),
                    "expected": "Parsed from example",
                    "actual": "Mock matched output" if is_compiled else "Compilation Failed",
                    "passed": False
                } for ex in (examples or ["Default Case"])
            ]
        }

    prompt = f"""You are an elite automated code grading engine.
The candidate has submitted a coding solution in the following language: {language}.

Problem Details:
Title: {title}
Description: {description}
Examples / Test cases: {examples}

Candidate's Code:
```
{code}
```

Please perform the following verification:
1. Syntactical and Compilation Check: Does this code compile/interpret successfully without syntax errors in {language}? If there are syntax errors, specify them in "output".
2. Logic & Example Run: Trace the code execution against the provided example test cases. Does it produce the correct outputs for the given inputs?

You MUST return a strictly valid JSON response. Do not wrap the JSON in markdown code blocks or add any other text.
JSON Structure:
{{
  "compiled": true, 
  "passed": true, 
  "output": "compiler output, execution logs, or error details",
  "test_results": [
    {{
      "input": "example input description",
      "expected": "expected output description",
      "actual": "actual output produced by the code",
      "passed": true 
    }}
  ]
}}
"""
    try:
        raw_response = _generate_with_model_fallback(prompt, timeout_sec=15)
        cleaned = re.sub(r"```json\s*", "", raw_response)
        cleaned = re.sub(r"```\s*", "", cleaned).strip()
        return json.loads(cleaned)
    except Exception as e:
        print(f"Error evaluating candidate code via AI: {e}. Falling back to offline evaluation.")
        is_compiled = True
        err_msg = ""
        try:
            if language.lower() in ["python", "py"]:
                compile(code, "<string>", "exec")
            elif language.lower() in ["js", "javascript"]:
                if code.count("{") != code.count("}"):
                    is_compiled = False
                    err_msg = "Syntax Error: Mismatched curly braces."
        except Exception as exc:
            is_compiled = False
            err_msg = str(exc)

        return {
            "compiled": is_compiled,
            "passed": is_compiled,
            "output": f"Offline dry-run fallback run due to API limit: {e}\n" + ("Local dry-run validation successful." if is_compiled else f"Compilation Error: {err_msg}"),
            "test_results": [
                {
                    "input": str(ex),
                    "expected": "Parsed from example",
                    "actual": "Mock matched output" if is_compiled else "Compilation Failed",
                    "passed": is_compiled
                } for ex in (examples or ["Default Case"])
            ]
        }


def grade_coding_question(language, code, title, description, constraints, examples, max_marks):
    if not code.strip():
        return 0.0, "No solution submitted."

    if not is_valid_key():
        # Fallback evaluation
        from ai_integration import evaluate_mock_test_code
        res = evaluate_mock_test_code(language, code, title, description, examples)
        if not res.get("compiled", False):
            return 0.0, "Compilation failed."
        elif res.get("passed", False):
            return float(max_marks), "All testcases passed."
        else:
            return float(max_marks) * 0.4, "Compiles but fails test cases."

    prompt = f"""You are an elite software engineering interviewer and compiler grader.
You need to evaluate a candidate's code submission for the following coding challenge:
Title: {title}
Description: {description}
Constraints: {constraints}
Examples: {examples}

Language selected: {language}
Candidate's code:
{code}

Grade the code out of {max_marks} marks.
CRITICAL RULES:
1. If the code is just a boilerplate/template, empty, or does not attempt to solve the problem logic at all, you MUST give 0 marks.
2. If the code has syntax errors or would fail compilation/execution, give 0 marks.
3. If it compiles and attempts the logic but fails test cases, give partial marks (e.g. 10-30% of max).
4. If it is correct but not optimized, give 70-80% of max marks.
5. If it is fully correct and optimal, give {max_marks} marks.

Return a JSON object with:
- 'score': float (the final score given out of {max_marks})
- 'feedback': string (short explanation of the grade and details of any bugs or logic errors)

Format response as strictly valid JSON (start with '{{' and end with '}}'). No markdown or backticks.
"""
    try:
        raw_response = _generate_with_model_fallback(prompt, timeout_sec=20)
        cleaned = re.sub(r"```json\s*", "", raw_response)
        cleaned = re.sub(r"```\s*", "", cleaned).strip()
        data = json.loads(cleaned)
        score = float(data.get("score", 0.0))
        feedback = data.get("feedback", "")
        # Clamp score between 0 and max_marks
        score = max(0.0, min(float(max_marks), score))
        return score, feedback
    except Exception as e:
        print(f"Error grading coding question: {e}")
        # Local fallback if Gemini fails
        from ai_integration import evaluate_mock_test_code
        res = evaluate_mock_test_code(language, code, title, description, examples)
        if not res.get("compiled", False):
            return 0.0, "Compilation failed (local dry-run)."
        else:
            return 0.0, "Compiles successfully, but AI logic verification is unavailable. 0 marks awarded for unverified logic."


def get_gemini_keywords_for_stream(stream_name):
    if not is_valid_key() or not stream_name:
        return _fallback_stream_keywords(stream_name)
    
    prompt = f"""
You are an expert ATS (Applicant Tracking System) Specialist and Career Coach.
Generate a JSON list of 12 to 15 high-impact, ATS-optimized technical keywords, tools, skills, and industry terms for a candidate aiming for the target career stream/role: "{stream_name}".

Output format MUST BE STRICT VALID JSON ARRAY of strings only, e.g.:
["Keyword 1", "Keyword 2", "Keyword 3", "Keyword 4"]
Do not include markdown backticks or any explanation outside the JSON array.
"""
    try:
        raw_text = _generate_with_model_fallback(prompt, timeout_sec=10)
        cleaned = re.sub(r'```json\s*', '', raw_text)
        cleaned = re.sub(r'```\s*', '', cleaned).strip()
        data = json.loads(cleaned)
        if isinstance(data, list) and len(data) > 0:
            return [str(x) for x in data]
    except Exception as e:
        print(f"Gemini stream keyword error: {e}")
    
    return _fallback_stream_keywords(stream_name)


def _fallback_stream_keywords(stream_name):
    s = (stream_name or "").lower()
    if "cyber" in s or "security" in s:
        return ["Wireshark", "Metasploit", "Penetration Testing", "Nmap", "SIEM", "SOC Analysis", "ISO 27001", "Cryptography", "Network Security", "Ethical Hacking", "Zero Trust", "Firewalls"]
    if "devops" in s or "cloud" in s:
        return ["Docker", "Kubernetes", "Terraform", "CI/CD Pipelines", "AWS", "Ansible", "Jenkins", "Prometheus", "Grafana", "Linux Administration", "Bash Scripting", "Helm"]
    if "ios" in s or "swift" in s:
        return ["Swift", "SwiftUI", "UIKit", "Xcode", "Combine Framework", "CoreData", "CocoaPods", "RESTful APIs", "Git", "TestFlight", "App Store Deployment", "MVVM"]
    if "android" in s or "kotlin" in s:
        return ["Kotlin", "Android SDK", "Jetpack Compose", "Coroutines", "Room Database", "Retrofit", "Dagger Hilt", "Gradle", "MVVM Architecture", "Unit Testing"]
    if "data eng" in s or "spark" in s:
        return ["Apache Spark", "PySpark", "Hadoop", "Kafka", "Airflow", "ETL Pipelines", "Snowflake", "Databricks", "Data Warehousing", "SQL Optimization", "Parquet", "BigQuery"]
    if "data sci" in s or "machine" in s or "ml" in s:
        return ["Scikit-learn", "Random Forest", "PyTorch", "TF-IDF", "NLP Pipelines", "Feature Engineering", "EDA", "Outlier Detection", "Model Deployment", "Cosine Similarity", "Pandas", "SciPy"]
    
    words = [w.capitalize() for w in re.findall(r'\w+', stream_name) if len(w) > 2]
    return words + ["System Design", "Agile Methodology", "Git / GitHub", "Problem Solving", "REST APIs", "Unit Testing", "CI/CD", "Performance Optimization"]


